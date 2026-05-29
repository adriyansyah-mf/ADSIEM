import io
import uuid
from typing import Annotated

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.deps import get_current_user, get_db, get_scoped_group
from app.models.models import SopDocument, User
from app.schemas.schemas import SopDocumentOut

router = APIRouter(tags=["sop"])

ALLOWED_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}
MAX_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


def _extract_text(content: bytes, content_type: str) -> str:
    if content_type == "application/pdf":
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        import docx
        doc = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return content.decode("utf-8", errors="replace")


@router.get("/api/sop-documents", response_model=list[SopDocumentOut])
async def list_sop_documents(
    db: Annotated[AsyncSession, Depends(get_db)],
    group_filter: Annotated[str | None, Depends(get_scoped_group)],
    _: Annotated[User, Depends(get_current_user)],
):
    q = select(SopDocument).order_by(SopDocument.created_at.desc())
    if group_filter:
        q = q.where(SopDocument.group_id == group_filter)
    rows = (await db.execute(q)).scalars().all()
    return rows


@router.post("/api/sop-documents", response_model=SopDocumentOut, status_code=201)
async def upload_sop_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    group_filter: str | None = Depends(get_scoped_group),
    current_user: User = Depends(get_current_user),
):
    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type: {file.content_type}. Use PDF, DOCX, or TXT.",
        )
    content = await file.read()
    if len(content) > MAX_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="File too large. Max 10 MB.")
    try:
        raw_text = _extract_text(content, file.content_type)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to extract text: {e}")
    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="Document contains no extractable text.")

    # Use the user's own group_id for scoping (superadmin uploads go to their group_id)
    effective_group = group_filter if group_filter else current_user.group_id

    doc = SopDocument(
        group_id=effective_group,
        filename=file.filename or "document",
        content_type=file.content_type,
        raw_text=raw_text,
        status="pending",
        uploaded_by=current_user.id,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc


@router.delete("/api/sop-documents/{doc_id}", status_code=204)
async def delete_sop_document(
    doc_id: uuid.UUID,
    db: Annotated[AsyncSession, Depends(get_db)],
    group_filter: Annotated[str | None, Depends(get_scoped_group)],
    _: Annotated[User, Depends(get_current_user)],
):
    doc = await db.get(SopDocument, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="SOP document not found")
    # Superadmin (group_filter=None) can delete any doc; others are scoped to their group
    if group_filter and doc.group_id != group_filter:
        raise HTTPException(status_code=404, detail="SOP document not found")
    await db.delete(doc)
    await db.commit()
